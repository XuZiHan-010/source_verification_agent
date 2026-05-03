"""Input document parsing for the source verification pipeline."""

from __future__ import annotations

import hashlib
import re
import tempfile
from pathlib import Path
from typing import Literal

from .errors import IngestError
from .schema import Block, IR

Format = Literal["pdf", "docx", "txt", "md", "auto"]


# pdfplumber 表格抽取参数：默认参数对中文带合并单元格的表格容易把一个逻辑 cell
# 切成多个相邻"碎片格"。优先用可见线条划列，避免文字间距被误判为列边界。
_PDF_TABLE_SETTINGS = {
    "vertical_strategy": "lines",
    "horizontal_strategy": "lines",
    "snap_tolerance": 4,
    "join_tolerance": 4,
    "edge_min_length": 12,
    "intersection_tolerance": 4,
    "text_tolerance": 3,
    "text_x_tolerance": 3,
    "text_y_tolerance": 3,
}


def ingest(path_or_bytes: str | Path | bytes, fmt: Format = "auto", doc_id: str | None = None) -> IR:
    """Parse an input file or text payload into the intermediate representation."""

    if isinstance(path_or_bytes, bytes):
        if fmt in ("auto", "txt", "md"):
            text = path_or_bytes.decode("utf-8", errors="replace")
            return _parse_text(text, doc_id=doc_id or _hash_doc_id(path_or_bytes), fmt="txt")
        suffix = f".{fmt}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as fh:
            fh.write(path_or_bytes)
            tmp_path = Path(fh.name)
        try:
            return ingest(tmp_path, fmt=fmt, doc_id=doc_id)
        finally:
            tmp_path.unlink(missing_ok=True)

    path = Path(path_or_bytes)
    if path.exists():
        actual_fmt = _detect_format(path, fmt)
        actual_doc_id = doc_id or _slug_doc_id(path.stem)
        if actual_fmt == "pdf":
            return _parse_pdf(path, actual_doc_id)
        if actual_fmt == "docx":
            return _parse_docx(path, actual_doc_id)
        if actual_fmt in ("txt", "md"):
            return _parse_text(path.read_text(encoding="utf-8", errors="replace"), actual_doc_id, actual_fmt)
        raise IngestError(f"Unsupported input format: {actual_fmt}", doc_id=actual_doc_id)

    if fmt in ("auto", "txt", "md"):
        text_fmt = "md" if fmt == "md" else "txt"
        return _parse_text(str(path_or_bytes), doc_id=doc_id or _hash_doc_id(str(path_or_bytes).encode()), fmt=text_fmt)
    raise IngestError(f"Input path does not exist: {path}")


def _detect_format(path: Path, fmt: Format) -> Literal["pdf", "docx", "txt", "md"]:
    if fmt != "auto":
        return fmt  # type: ignore[return-value]
    suffix = path.suffix.lower().lstrip(".")
    if suffix in {"pdf", "docx", "txt", "md"}:
        return suffix  # type: ignore[return-value]
    raise IngestError(f"Cannot detect input format from suffix: {path.suffix}", doc_id=path.stem)


def _parse_pdf(path: Path, doc_id: str) -> IR:
    try:
        import fitz
        import pdfplumber
    except ImportError as exc:
        raise IngestError("PDF parsing requires pdfplumber and pymupdf", doc_id=doc_id) from exc

    blocks: list[Block] = []
    try:
        with fitz.open(path) as fitz_doc, pdfplumber.open(path) as pdf:
            previous_table: Block | None = None
            previous_header: list[str] | None = None
            for page_idx, page in enumerate(pdf.pages, start=1):
                fitz_page = fitz_doc[page_idx - 1]
                page_links = _extract_page_links(fitz_page)
                footnotes = _extract_page_footnotes(fitz_page, page)
                tables_lines = page.extract_tables(table_settings=_PDF_TABLE_SETTINGS) or []
                tables = [table for table in tables_lines if _is_main_table(table)]
                if not tables:
                    # fallback：无明显边框的表回退到 pdfplumber 默认（text 策略）
                    tables = [table for table in page.extract_tables() or [] if _is_main_table(table)]
                if tables:
                    for table_idx, rows in enumerate(tables):
                        clean_rows = _fill_table(rows)
                        header = _header_signature(clean_rows)
                        if (
                            previous_table
                            and previous_table.rows is not None
                            and previous_header
                            and header
                            and _headers_match(previous_header, header)
                        ):
                            clean_rows = clean_rows[1:]
                            if clean_rows:
                                previous_table.rows.extend(clean_rows)
                                previous_table.footnotes.update(footnotes)
                            continue
                        if (
                            previous_table
                            and previous_table.rows is not None
                            and previous_header
                            and not header
                            and _is_table_continuation(clean_rows)
                        ):
                            previous_table.rows.extend(clean_rows)
                            previous_table.footnotes.update(footnotes)
                            continue
                        link_map = {
                            f"{row},{col}": url
                            for (row, col), url in page_links.get(table_idx, {}).items()
                        }
                        block = Block(
                            type="table",
                            page=page_idx,
                            rows=clean_rows,
                            hyperlinks=link_map,
                            footnotes=footnotes,
                        )
                        blocks.append(block)
                        previous_table = block
                        previous_header = header
                text = page.extract_text() or ""
                blocks.extend(_text_to_blocks(text, page=page_idx))
    except Exception as exc:  # pragma: no cover - library-specific failures
        raise IngestError(str(exc), doc_id=doc_id) from exc
    return IR(doc_id=doc_id, source_format="pdf", blocks=blocks)


def _pdf_links(path: Path, fitz_module) -> dict[int, dict[int, dict[tuple[int, int], str]]]:
    links: dict[int, dict[int, dict[tuple[int, int], str]]] = {}
    try:
        with fitz_module.open(path) as doc:
            for page_idx, page in enumerate(doc, start=1):
                urls = [link.get("uri") for link in page.get_links() if link.get("uri")]
                if urls:
                    links[page_idx] = {0: {(0, idx): url for idx, url in enumerate(urls)}}
    except Exception:
        return {}
    return links


def _extract_page_links(page) -> dict[int, dict[tuple[int, int], str]]:
    urls = [link.get("uri") for link in page.get_links() if link.get("uri")]
    return {0: {(0, idx): url for idx, url in enumerate(urls)}} if urls else {}


def _extract_page_footnotes(page_fitz, page_plumber) -> dict[int, str]:
    footnotes: dict[int, str] = {}
    for link in page_fitz.get_links():
        url = link.get("uri")
        rect = link.get("from")
        if not url or rect is None:
            continue
        number = _nearest_link_number(page_fitz, rect)
        if number is not None:
            footnotes[number] = url

    bottom_text = page_plumber.extract_text() or ""
    current_number: int | None = None
    current_url = ""
    for line in bottom_text.splitlines():
        stripped = line.strip()
        match = re.match(r"^(\d{1,3})\s*(https?://\S+)", stripped)
        if match:
            if current_number is not None and current_url:
                footnotes[current_number] = current_url.rstrip(").,;，。")
            current_number = int(match.group(1))
            current_url = match.group(2)
        elif current_number is not None and stripped and re.match(r"^[A-Za-z0-9_./?=&%#:+-]+$", stripped):
            current_url += stripped
    if current_number is not None and current_url:
        footnotes[current_number] = current_url.rstrip(").,;，。")
    return footnotes


def _nearest_link_number(page, rect) -> int | None:
    try:
        words = page.get_text("words")
    except Exception:
        return None
    candidates: list[tuple[float, int]] = []
    for x0, y0, x1, y1, word, *_ in words:
        if not re.fullmatch(r"\d{1,3}", str(word)):
            continue
        if y1 < rect.y0 - 8 or y0 > rect.y1 + 8:
            continue
        if x0 > rect.x0 + 4:
            continue
        distance = abs(y0 - rect.y0) + max(0.0, rect.x0 - x1)
        candidates.append((distance, int(word)))
    if not candidates:
        return None
    return min(candidates, key=lambda item: item[0])[1]


HEADER_KEYWORDS = (
    "指标",
    "数值",
    "年份",
    "来源",
    "备注",
    "驱动因素",
    "原文表述",
    "摘要",
    "差异说明",
    "缺失信息",
    "当前状态",
    "source",
    "metric",
    "value",
    "year",
    "政策",
    "法规",
    "标准",
    "发布机构",
    "发布时间",
    "适用",
    "核心内容",
    "监管事项",
    "具体要求",
    "适用对象",
)


def _header_signature(rows: list[list[str | None]]) -> list[str] | None:
    if not rows:
        return None
    cells = [_normalise_header_cell(cell) for cell in rows[0]]
    hits = sum(1 for cell in cells for keyword in HEADER_KEYWORDS if keyword.lower() in cell.lower())
    return cells if hits >= 2 else None


def _headers_match(left: list[str], right: list[str]) -> bool:
    if not left or not right:
        return False
    overlap = sum(1 for a, b in zip(left, right) if a and b and a == b)
    return overlap >= max(2, min(len(left), len(right)) // 2)


def _normalise_header_cell(value: str | None) -> str:
    return re.sub(r"\s+", "", value or "").strip()


def _is_main_table(rows: list[list[str | None]]) -> bool:
    return any(len(row) >= 4 for row in rows)


def _is_table_continuation(rows: list[list[str | None]]) -> bool:
    if not rows:
        return False
    return any(len(row) >= 4 for row in rows) and _header_signature(rows) is None


def _parse_docx(path: Path, doc_id: str) -> IR:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestError("DOCX parsing requires python-docx", doc_id=doc_id) from exc

    blocks: list[Block] = []
    try:
        doc = Document(path)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name.lower() if paragraph.style else ""
            if "heading" in style:
                level_match = re.search(r"(\d+)", style)
                level = int(level_match.group(1)) if level_match else 1
                blocks.append(Block(type="heading", level=level, text=text))
            else:
                blocks.append(Block(type="paragraph", text=text))
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            blocks.append(Block(type="table", rows=_fill_table(rows)))
    except Exception as exc:  # pragma: no cover - library-specific failures
        raise IngestError(str(exc), doc_id=doc_id) from exc
    return IR(doc_id=doc_id, source_format="docx", blocks=blocks)


def _parse_text(text: str, doc_id: str, fmt: Literal["txt", "md"] = "txt") -> IR:
    blocks = _parse_markdown_tables(text) if fmt == "md" else []
    if not blocks:
        blocks = _text_to_blocks(text)
    return IR(doc_id=doc_id, source_format=fmt, blocks=blocks)


def _text_to_blocks(text: str, page: int | None = None) -> list[Block]:
    blocks: list[Block] = []
    for raw in re.split(r"\n\s*\n", text):
        chunk = raw.strip()
        if not chunk:
            continue
        if _looks_like_table(chunk):
            rows = [[cell.strip() for cell in re.split(r"\t| {2,}|\|", line) if cell.strip()] for line in chunk.splitlines()]
            if rows:
                blocks.append(Block(type="table", page=page, rows=_fill_table(rows)))
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", chunk)
        if heading:
            blocks.append(Block(type="heading", level=len(heading.group(1)), page=page, text=heading.group(2).strip()))
        elif len(chunk) <= 80 and re.match(r"^(\d+(\.\d+)*[、.\s])?[\w\u4e00-\u9fff].*$", chunk):
            blocks.append(Block(type="heading", level=2, page=page, text=chunk))
        else:
            blocks.append(Block(type="paragraph", page=page, text=chunk))
    return blocks


def _parse_markdown_tables(text: str) -> list[Block]:
    blocks: list[Block] = []
    current_paragraph: list[str] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("#"):
            if current_paragraph:
                blocks.append(Block(type="paragraph", text="\n".join(current_paragraph).strip()))
                current_paragraph = []
            level = len(line) - len(line.lstrip("#"))
            blocks.append(Block(type="heading", level=level, text=line[level:].strip()))
            i += 1
            continue
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:\-|]+\|", lines[i + 1]):
            if current_paragraph:
                blocks.append(Block(type="paragraph", text="\n".join(current_paragraph).strip()))
                current_paragraph = []
            rows = []
            while i < len(lines) and "|" in lines[i]:
                row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not _is_markdown_separator(row):
                    rows.append(row)
                i += 1
            blocks.append(Block(type="table", rows=_fill_table(rows)))
            continue
        if line:
            current_paragraph.append(line)
        i += 1
    if current_paragraph:
        blocks.append(Block(type="paragraph", text="\n".join(current_paragraph).strip()))
    return blocks


def _fill_table(rows: list[list[str | None]]) -> list[list[str | None]]:
    filled: list[list[str | None]] = []
    previous: list[str | None] = []
    for row in rows:
        raw = [cell.strip() if isinstance(cell, str) else cell for cell in row]
        if filled and _is_continuation_row(raw):
            target = filled[-1]
            for idx, value in enumerate(raw):
                if value in ("", None):
                    continue
                while idx >= len(target):
                    target.append(None)
                target[idx] = f"{target[idx]} {value}".strip() if target[idx] else value
            previous = target
            continue

        out: list[str | None] = []
        for idx, cell in enumerate(raw):
            value = cell.strip() if isinstance(cell, str) else cell
            if value in ("", None) and idx < len(previous):
                value = previous[idx]
            out.append(value)
        previous = out
        filled.append(out)
    return filled


def _is_continuation_row(row: list[str | None]) -> bool:
    if not any(cell for cell in row):
        return False
    if len(row) < 5:
        return False
    leading = row[:5]
    if all(cell in ("", None) for cell in leading):
        return True
    # 放宽：前 5 列大部分空，且非空内容是短碎片（典型 PDF 行内换行）
    leading_empty = sum(1 for cell in leading if cell in ("", None))
    nonempty = [cell for cell in row if cell]
    if leading_empty >= 4 and nonempty and all(len(re.sub(r"\s+", "", cell)) <= 6 for cell in nonempty):
        return True
    return False


def _looks_like_table(text: str) -> bool:
    lines = [line for line in text.splitlines() if line.strip()]
    return len(lines) >= 2 and sum(1 for line in lines if "\t" in line or "|" in line or re.search(r" {2,}", line)) >= 2


def _is_markdown_separator(row: list[str]) -> bool:
    return bool(row) and all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell or "") for cell in row)


def _slug_doc_id(value: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", value, flags=re.UNICODE).strip("_")
    return slug or "document"


def _hash_doc_id(value: bytes) -> str:
    return hashlib.sha1(value).hexdigest()[:12]
